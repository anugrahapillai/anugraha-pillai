from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = "/Users/ratiraj/project/Anughara/Project_Transfer_Document_Anugraha_Pillai.docx"
BLUE = "244C66"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = RGBColor(92, 101, 110)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_key_value_table(doc, rows, widths=(2700, 6660)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_GRAY)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(BLUE)
    set_table_geometry(table, list(widths))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_link_row(table, label, url):
    cells = table.add_row().cells
    cells[0].text = label
    set_cell_shading(cells[0], LIGHT_GRAY)
    for run in cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
    p = cells[1].paragraphs[0]
    add_hyperlink(p, url, url)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.45)
section.footer_distance = Inches(0.45)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(30, 38, 45)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
for style_name, size, before, after in (("Heading 1", 16, 16, 7), ("Heading 2", 13, 11, 5), ("Heading 3", 11.5, 8, 4)):
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(BLUE)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for style_name in ("List Bullet", "List Number"):
    st = styles[style_name]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.38)
    st.paragraph_format.first_line_indent = Inches(-0.19)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.15

header = section.header.paragraphs[0]
header.text = "ANUGRAHA PILLAI WEBSITE  |  PROJECT HANDOVER"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in header.runs:
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = MUTED

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Confidential • Store securely • Do not share publicly")
run.font.name = "Calibri"
run.font.size = Pt(8)
run.font.color.rgb = MUTED

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("PROJECT TRANSFER & HANDOVER DOCUMENT")
r.font.name = "Calibri"
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(BLUE)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
r = p.add_run("Anugraha Pillai Portfolio Website")
r.font.name = "Calibri"
r.font.size = Pt(14)
r.font.color.rgb = MUTED

meta = add_key_value_table(doc, [
    ("Website", "www.anugrahapillai.in"),
    ("Prepared by", "Ratiraj Chavan"),
    ("Transfer date", "16 August 2026"),
    ("Production launch date", "To be confirmed"),
])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("CONFIDENTIAL ACCESS INFORMATION")
r.bold = True
r.font.color.rgb = RGBColor(139, 69, 19)
p.add_run("  This document contains passwords. Keep it in a secure location and change every shared password after the handover is accepted.")

doc.add_heading("1. Purpose of this document", level=1)
doc.add_paragraph("This document transfers the website links, account access details, maintenance arrangement and commercial terms for the Anugraha Pillai website.")

doc.add_heading("2. Website and administration", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.rows[0].cells[0].text = "Item"
table.rows[0].cells[1].text = "Link"
for cell in table.rows[0].cells:
    set_cell_shading(cell, LIGHT_BLUE)
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
set_repeat_table_header(table.rows[0])
add_link_row(table, "Live website", "https://www.anugrahapillai.in/")
add_link_row(table, "Admin CMS", "https://www.anugrahapillai.in/admin/login")
set_table_geometry(table, [2700, 6660])

doc.add_heading("3. Account access", level=1)
doc.add_heading("3.1 Primary Google account", level=2)
add_key_value_table(doc, [
    ("Email", "anugrahapillai.work@gmail.com"),
    ("Password", "workanugrahapillai00"),
    ("Connected services", "GitHub, Firebase and Vercel use Google sign-in."),
])

doc.add_heading("3.2 EmailJS service", level=2)
emailjs = add_key_value_table(doc, [
    ("Sign-in URL", "https://dashboard.emailjs.com/sign-in"),
    ("Email", "anugrahapillai.work@gmail.com"),
    ("Password", "Anugraha@123"),
])

doc.add_heading("3.3 Admin CMS accounts", level=2)
cms = doc.add_table(rows=1, cols=3)
cms.style = "Table Grid"
headers = ["Account", "Email", "Password"]
for idx, text in enumerate(headers):
    cms.rows[0].cells[idx].text = text
    set_cell_shading(cms.rows[0].cells[idx], LIGHT_BLUE)
    for run in cms.rows[0].cells[idx].paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
for row in [
    ("Admin 1", "anugrahaapillai@gmail.com", "anugrahap01"),
    ("Admin 2", "anugraha.twinkle@gmail.com", "anugraha@1234"),
]:
    cells = cms.add_row().cells
    for idx, value in enumerate(row):
        cells[idx].text = value
set_repeat_table_header(cms.rows[0])
set_table_geometry(cms, [1500, 4800, 3060])
p = doc.add_paragraph("CMS passwords can be changed using the “Forgot password” option on the admin login page.")
p.paragraph_format.space_before = Pt(5)

doc.add_heading("4. Connected platforms", level=1)
platforms = doc.add_table(rows=1, cols=2)
platforms.style = "Table Grid"
platforms.rows[0].cells[0].text = "Platform"
platforms.rows[0].cells[1].text = "Project link"
for cell in platforms.rows[0].cells:
    set_cell_shading(cell, LIGHT_BLUE)
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
set_repeat_table_header(platforms.rows[0])
for label, url in [
    ("GitHub", "https://github.com/anugrahapillai/anugraha-pillai"),
    ("Firebase", "https://console.firebase.google.com/project/anugraha-pillai/overview"),
    ("Vercel", "https://vercel.com/anugraha-pillai/anugraha-pillai"),
]:
    add_link_row(platforms, label, url)
set_table_geometry(platforms, [2700, 6660])
doc.add_paragraph("Access these services with the primary Google account listed in Section 3.1.")

doc.add_heading("5. Search indexing and SEO", level=1)
doc.add_paragraph("Google indexing and basic SEO setup have been added to the website. Future SEO, metadata and indexing changes can be made through the relevant website, Google and platform settings.")

doc.add_heading("6. Handover checklist", level=1)
for item in [
    "Confirm that the live website opens correctly.",
    "Confirm access to the Admin CMS with both accounts.",
    "Confirm access to the primary Google account.",
    "Confirm access to EmailJS, GitHub, Firebase and Vercel.",
    "Change all shared passwords and enable two-factor authentication where available.",
    "Store the updated credentials in a secure password manager.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("7. Annual Maintenance Contract (AMC)", level=1)
doc.add_paragraph("A basic Annual Maintenance Contract (AMC) is included at no additional charge for 12 months from the production launch date. It keeps the approved website operational; it is not a redesign or enhancement allowance.")

doc.add_heading("7.1 Included support", level=2)
for item in [
    "One scheduled maintenance meeting and audit each month for 12 months.",
    "Basic bug fixes within the approved and delivered website scope.",
    "Basic checks of website availability and core functions.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("7.2 Not included", level=2)
for item in [
    "Redesign, new layouts, colour changes or typography changes.",
    "New features, integrations, pages or major content changes.",
    "Third-party charges, hosting upgrades, storage, bandwidth or premium services.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("8. Monthly maintenance meetings", level=1)
doc.add_paragraph("Twelve meetings will be held, normally one meeting each month. Dates will be agreed by both parties. A missed or postponed meeting should be rescheduled within the same month where reasonably possible.")
meetings = doc.add_table(rows=1, cols=3)
meetings.style = "Table Grid"
for idx, text in enumerate(("Meeting", "Target period", "Status / notes")):
    meetings.rows[0].cells[idx].text = text
    set_cell_shading(meetings.rows[0].cells[idx], LIGHT_BLUE)
    for run in meetings.rows[0].cells[idx].paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
for i in range(1, 13):
    cells = meetings.add_row().cells
    cells[0].text = str(i)
    cells[1].text = f"Month {i}"
    cells[2].text = ""
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
set_repeat_table_header(meetings.rows[0])
set_table_geometry(meetings, [1300, 2800, 5260])

doc.add_heading("9. Suggested payment schedule", level=1)
payments = doc.add_table(rows=1, cols=3)
payments.style = "Table Grid"
for idx, text in enumerate(("Milestone", "Payment", "Due")):
    payments.rows[0].cells[idx].text = text
    set_cell_shading(payments.rows[0].cells[idx], LIGHT_BLUE)
    for run in payments.rows[0].cells[idx].paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
for row in [
    ("Project commencement", "40%", "Before work begins; payment must be received and cleared."),
    ("Design and development approval", "30%", "After approval of the main design and functional build."),
    ("Production launch and handover", "30%", "Before final launch, credential transfer and handover."),
]:
    cells = payments.add_row().cells
    for idx, value in enumerate(row):
        cells[idx].text = value
set_repeat_table_header(payments.rows[0])
set_table_geometry(payments, [3000, 1300, 5060])

doc.add_heading("10. Terms and conditions", level=1)
terms = [
    "All payments made to Ratiraj Chavan are non-refundable, including the initial 40% payment and subsequent milestone payments.",
    "The one-month delivery period begins only when the initial 40% payment has been received and cleared. Client delays in supplying content, approvals, access credentials or consolidated feedback extend the delivery date by the corresponding delay.",
    "Extra Firebase or other cloud storage, bandwidth, database, hosting, email, premium-domain or third-party service charges are not the responsibility of Ratiraj Chavan and must be paid by the client.",
    "The included domain obligation is limited to registering one approved standard .in domain for three years. Ratiraj Chavan has no responsibility for renewal, expiry, recovery or related charges after that three-year term.",
    "The one-year AMC includes one monthly audit and basic bug fixes within the delivered scope. Layout, colour, typography, visual-design, content and feature changes remain excluded and require a separate quotation.",
]
for item in terms:
    doc.add_paragraph(item, style="List Number")

doc.add_page_break()
accept_heading = doc.add_heading("11. Acceptance", level=1)
keep_with_next(accept_heading)
accept_text = doc.add_paragraph("By signing below, both parties confirm that the listed access information and project assets have been handed over, subject to the terms in this document.")
keep_with_next(accept_text)
sign = doc.add_table(rows=3, cols=2)
sign.style = "Table Grid"
for idx, heading in enumerate(("Client", "Service provider")):
    sign.rows[0].cells[idx].text = heading
    set_cell_shading(sign.rows[0].cells[idx], LIGHT_BLUE)
    for run in sign.rows[0].cells[idx].paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
sign.rows[1].cells[0].text = "Name and signature:\n\n"
sign.rows[1].cells[1].text = "Ratiraj Chavan\nSignature:\n"
sign.rows[2].cells[0].text = "Date:"
sign.rows[2].cells[1].text = "Date:"
set_table_geometry(sign, [4680, 4680])

doc.core_properties.title = "Project Transfer and Handover Document — Anugraha Pillai"
doc.core_properties.subject = "Website access, AMC, payment schedule and terms"
doc.core_properties.author = "Ratiraj Chavan"
doc.core_properties.keywords = "handover, website, AMC, credentials, Anugraha Pillai"
doc.save(OUT)
print(OUT)
